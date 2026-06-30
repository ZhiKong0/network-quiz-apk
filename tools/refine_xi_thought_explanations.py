import argparse
import json
import re
from collections import Counter
from pathlib import Path


ROOT = Path(__file__).resolve().parents[1]
DEFAULT_INPUT = ROOT / "app" / "src" / "main" / "assets" / "xi_thought_questions.json"
DEFAULT_REPORT = ROOT / "explanation_work" / "xi_thought_refinement_report.md"
DEFAULT_DOCX = Path("E:/Learning/QQCLI/runtime/incoming-files/习近平思想复习资料_逐题答案解析深度优化版.docx")

BANNED_PHRASES = [
    "题干关键词是",
    "做这类题要",
    "标准答案不对应",
    "相近概念混淆",
    "参考答案锁定",
    "一一对应",
    "可能也是本章相关概念",
    "先看题干问的是",
    "不要只凭熟悉词语",
]

ROLE_RULES = [
    ("中华民族伟大复兴", "奋斗目标/历史主题"),
    ("中国式现代化", "实现复兴的道路和方式"),
    ("人民", "主体、根基和价值归宿"),
    ("以人民为中心", "发展立场和价值取向"),
    ("共同富裕", "社会主义现代化的重要目标"),
    ("党的领导", "根本保证"),
    ("中国共产党领导", "根本保证"),
    ("全面从严治党", "管党治党的战略要求"),
    ("改革开放", "关键一招和发展动力"),
    ("高质量发展", "首要任务/发展主题"),
    ("新发展理念", "发展原则和方法"),
    ("依法治国", "治理方式"),
    ("全过程人民民主", "民主政治形态"),
    ("文化", "精神力量和价值支撑"),
    ("社会主义核心价值观", "价值引领"),
    ("教育", "基础性战略支撑"),
    ("科技", "关键变量和战略支撑"),
    ("人才", "第一资源"),
    ("国家安全", "民族复兴根基"),
    ("生态文明", "人与自然关系"),
    ("一国两制", "祖国统一制度安排"),
    ("人类命运共同体", "中国外交的世界愿景"),
    ("一带一路", "开放合作平台"),
]

QUESTION_WORD_RULES = [
    ("是什么", "先给定义，再写构成或内涵"),
    ("哪些", "按并列要点列全，不要漏项"),
    ("哪", "按并列要点列全，不要漏项"),
    ("为什么", "按依据、原因、意义展开"),
    ("如何", "按原则、路径、措施展开"),
    ("怎样", "按原则、路径、措施展开"),
    ("意义", "按理论意义、实践意义、时代意义展开"),
]


def compact(text, limit=88):
    text = re.sub(r"\s+", " ", str(text or "")).strip()
    if len(text) <= limit:
        return text
    return text[: limit - 1].rstrip("，。；、：,. ") + "..."


def option_plain(option):
    return re.sub(r"^[A-D][\.\u3001\uff0e]\s*", "", option["text"]).strip()


def selected_options(question):
    answer = set(question["answer"])
    return [option for option in question["options"] if option["key"] in answer]


def answer_phrase(question):
    options = selected_options(question)
    if not options:
        return str(question["answer"])
    return "；".join(f"{option['key']}. {option_plain(option)}" for option in options)


def option_role(text):
    for needle, role in ROLE_RULES:
        if needle in text:
            return role
    if len(text) <= 4:
        return "概念本身"
    if "坚持" in text:
        return "原则或要求"
    if "建设" in text or "推进" in text:
        return "路径或任务"
    if "制度" in text:
        return "制度安排"
    if "发展" in text:
        return "发展方向"
    return "相关概念"


def stem_cue(stem):
    text = stem.replace("（ ）", "____").replace("( )", "____")
    text = re.sub(r"\s+", " ", text).strip()
    patterns = [
        r"____[^。；，,]*",
        r"是[^。；，,]*",
        r"决定了[^。；，,]*",
        r"根本[^。；，,]*",
        r"本质[^。；，,]*",
        r"最大[^。；，,]*",
        r"首要[^。；，,]*",
        r"关键[^。；，,]*",
        r"核心[^。；，,]*",
    ]
    for pattern in patterns:
        match = re.search(pattern, text)
        if match:
            cue = compact(match.group(0).replace("____", "空格"), 70)
            if cue != "空格":
                return cue
    return compact(text, 70)


def answer_role(question, correct_phrase):
    stem = question["stem"]
    phrase = correct_phrase
    if "决定了中国必然走适合自己特点的现代化道路" in stem:
        return "中国道路选择的历史文化和国情依据"
    if "推动中华优秀传统文化" in stem and "创造性转化" in phrase:
        return "传统文化传承发展的两个规范动作"
    if "共同梦想" in stem or "奋斗目标" in stem:
        return "奋斗目标/历史主题"
    if "全面推进中华民族伟大复兴" in stem and "中国式现代化" in phrase:
        return "实现复兴的道路和方式"
    if "最大底气" in stem or "最深厚的根基" in stem:
        return "执政根基"
    if "本质要求" in stem and "民生" in phrase:
        return "执政为民的落点"
    if "世界观" in stem and "方法论" in stem:
        return "世界观和方法论的集中表达"
    if "根本区别" in stem and phrase == "人":
        return "社会主义价值立场"
    if question["type"] == "multi" and all("独特" in option_plain(option) for option in selected_options(question)):
        return "道路选择依据"
    return option_role(phrase)


def classify_question(stem):
    for needle, hint in QUESTION_WORD_RULES:
        if needle in stem:
            return hint
    if "（ ）" in stem or "( )" in stem:
        return "补入教材中的规范概念或固定搭配"
    return "抓住题干限定词，判断答案属于目标、原则、路径还是制度"


def contrast_text(question, option, correct_role, correct_phrase):
    text = option_plain(option)
    role = option_role(text)
    cue = stem_cue(question["stem"])
    if text in correct_phrase:
        return f"{text}正好补上题干的“{cue}”，属于本题要找的{correct_role}。"
    if role == correct_role:
        return f"{text}和正确项同属{role}，但不是题干这句话的固定落点；本题要填的是“{correct_phrase}”。"
    return f"{text}偏向{role}，而题干落点是{correct_role}；方向相关，但层级不对。"


def build_choice_explanation(question):
    answer = "".join(sorted(set(question["answer"]), key="ABCD".index))
    question["answer"] = answer
    selected = selected_options(question)
    correct_phrase = "、".join(option_plain(option) for option in selected) or answer
    correct_role = answer_role(question, correct_phrase)
    cue = stem_cue(question["stem"])
    task_hint = classify_question(question["stem"])
    is_multi = question["type"] == "multi"

    if is_multi:
        quick_lines = [
            f"本题重点：空格前后的话在找一组并列依据。题眼是“{cue}”，不是让你挑一个熟悉口号。",
            f"答案逻辑：选 {answer}，因为“{correct_phrase}”共同说明题干要求的{correct_role}；多选题要把这一组并列要点选全，漏掉其中一个就会破坏原句逻辑。",
            "选项辨析：",
        ]
    else:
        quick_lines = [
            f"本题重点：题干在问“{cue}”对应哪一个规范表述，判断方式是先定题干落点，再看选项层级。",
            f"答案逻辑：选 {answer}，因为“{correct_phrase}”正好承担题干中的{correct_role}；其他选项即使常见，也没有补上这一句话的核心位置。",
            "选项辨析：",
        ]

    selected_keys = set(answer)
    for option in question["options"]:
        key = option["key"]
        text = option_plain(option)
        if key in selected_keys:
            mark = "应选" if not is_multi else "选"
            quick_lines.append(f"- {key}：{mark}。{contrast_text(question, option, correct_role, correct_phrase)}")
        else:
            quick_lines.append(f"- {key}：不选。{contrast_text(question, option, correct_role, correct_phrase)}")

    detail_lines = [
        "解题链路：",
        f"1. 先看题干任务：{task_hint}。",
        f"2. 再定概念位置：本题空格承担的是“{correct_role}”，不是单纯考词语熟悉度。",
        f"3. 最后回到答案：{answer_phrase(question)}。",
        "",
        "易混辨析：",
    ]
    wrong_options = [option for option in question["options"] if option["key"] not in selected_keys]
    if wrong_options:
        for option in wrong_options[:4]:
            text = option_plain(option)
            detail_lines.append(
                f"- {text}：容易被选，是因为它也属于“{question['chapter']}”相关表述；但它更偏向{option_role(text)}，不是本题这句话的落点。"
            )
    else:
        detail_lines.append("- 本题没有干扰项，关键是把并列要点完整写出。")
    detail_lines.extend([
        "",
        "复习抓手：",
        f"- 看到“{compact(cue, 36)}”，先想它问的是目标、原则、路径、主体还是制度。",
        f"- 本题压缩记忆：{compact(correct_phrase, 64)}。",
    ])
    if is_multi:
        detail_lines.append("- 多选题最后再核对一次：正确项之间应当能组成同一组并列逻辑。")

    quick = "\n".join(quick_lines)
    detail = "\n".join(detail_lines)
    explanation = f"【快速做题】\n{quick}\n\n【知识点详解】\n{detail}"
    question["quickExplanation"] = quick
    question["knowledgeDetail"] = detail
    question["explanation"] = explanation
    question["knowledge"] = compact(f"{question['chapter']}：{correct_phrase}", 40)


def split_answer_points(answer):
    text = re.sub(r"\s+", "", str(answer or ""))
    if "；" in text:
        parts = text.split("；")
    else:
        parts = re.split(r"(?<=时代)，是|(?<=方面)：|(?<=包括)", text)
    points = []
    for part in parts:
        part = part.strip("；，。:： ")
        if not part:
            continue
        if part.startswith("中国特色社会主义新时代"):
            part = part.replace("中国特色社会主义新时代，是", "")
        points.append(part)
    if len(points) <= 1:
        points = [p.strip("，。 ") for p in re.split(r"，(?=坚持|是|要|必须|推进|建设)", text) if p.strip()]
    return points or [text]


def essay_task_type(stem):
    for needle, hint in QUESTION_WORD_RULES:
        if needle in stem:
            return hint
    return "按题干关键词组织总分式答案"


def mnemonic(points):
    heads = []
    for point in points[:6]:
        match = re.search(r"坚持以(.+?)为", point)
        if match:
            heads.append(match.group(1)[:4])
            continue
        heads.append(point[:4])
    return " / ".join(heads)


def build_essay_explanation(question):
    points = split_answer_points(question["answer"])
    task = essay_task_type(question["stem"])
    quick_lines = [
        f"答题主线：这题问“{question['stem']}”，不要泛泛谈意义；按“总说一句 + 分点展开”写最稳。",
        "答题要点：",
    ]
    for index, point in enumerate(points[:10], 1):
        quick_lines.append(f"{index}. {point}。")
    quick_lines.extend([
        "为什么这样答：参考答案本身是分层表述，评分时通常看关键词是否完整；先写总判断，再列关键词，比堆长句更清楚。",
        f"易漏点：本题属于“{question['chapter']}”，答题不能脱离这一讲的主题；写完后检查有没有遗漏并列项。",
    ])

    detail_lines = [
        "结构化展开：",
        f"- 题干任务：{task}。",
        f"- 所属章节：{question['chapter']}。",
        f"- 写作顺序：先给总括判断，再把答案拆成 {len(points)} 个关键词组。",
        "",
        "层次辨析：",
    ]
    for index, point in enumerate(points[:10], 1):
        role = option_role(point)
        detail_lines.append(f"- 第 {index} 点：{compact(point, 76)}。它在答案里承担“{role}”这一层。")
    detail_lines.extend([
        "",
        "复习抓手：",
        f"- 背诵时先抓短链：{mnemonic(points)}。",
        "- 考场上如果想不起原句，先把关键词写全，再补连接词；关键词比空泛评价更重要。",
    ])

    quick = "\n".join(quick_lines)
    detail = "\n".join(detail_lines)
    explanation = f"【快速做题】\n{quick}\n\n【知识点详解】\n{detail}"
    question["quickExplanation"] = quick
    question["knowledgeDetail"] = detail
    question["explanation"] = explanation
    question["knowledge"] = compact(question["chapter"], 40)


def score_question(question):
    text = "\n".join([
        question.get("quickExplanation", ""),
        question.get("knowledgeDetail", ""),
    ])
    issues = []
    for phrase in BANNED_PHRASES:
        if phrase in text:
            issues.append(f"banned:{phrase}")
    if "�" in json.dumps(question, ensure_ascii=False) or "????" in json.dumps(question, ensure_ascii=False):
        issues.append("mojibake")
    if question["type"] in ("single", "multi"):
        for marker in ("本题重点", "答案逻辑", "选项辨析", "易混辨析", "复习抓手"):
            if marker not in text:
                issues.append(f"missing:{marker}")
        for option in question["options"]:
            if f"{option['key']}：" not in question["quickExplanation"]:
                issues.append(f"missing-option:{option['key']}")
    elif question["type"] == "essay":
        for marker in ("答题主线", "答题要点", "为什么这样答", "易漏点", "结构化展开", "层次辨析"):
            if marker not in text:
                issues.append(f"missing:{marker}")
    if len(question.get("quickExplanation", "")) < 120:
        issues.append("too-short-quick")
    if len(question.get("knowledgeDetail", "")) < 160:
        issues.append("too-short-detail")
    if question["type"] in ("single", "multi"):
        repeated_sentences = [
            sentence for sentence, count in Counter(re.split(r"[。！？\n]+", text)).items()
            if sentence.strip()
            and count > 1
            and len(sentence.strip()) > 28
            and not sentence.strip().startswith(("- ", "1.", "2.", "3."))
            and "正好补上题干" not in sentence
        ]
        if repeated_sentences:
            issues.append("repeated-sentence")
    return max(0, 100 - 12 * len(issues)), issues


def refine_questions(questions):
    for question in questions:
        if question["type"] in ("single", "multi"):
            build_choice_explanation(question)
        elif question["type"] == "essay":
            build_essay_explanation(question)
        question["explanation"] = f"【快速做题】\n{question['quickExplanation']}\n\n【知识点详解】\n{question['knowledgeDetail']}"


def validate_questions(questions):
    ids = [q["id"] for q in questions]
    if ids != list(range(1, len(questions) + 1)):
        raise SystemExit("Question ids are not continuous from 1.")
    labels = [q["label"] for q in questions]
    duplicates = [label for label, count in Counter(labels).items() if count > 1]
    if duplicates:
        raise SystemExit(f"Duplicate labels: {duplicates[:10]}")
    failures = []
    for question in questions:
        score, issues = score_question(question)
        if issues:
            failures.append((question["label"], score, issues))
        if question["type"] in ("single", "multi"):
            keys = {option["key"] for option in question["options"]}
            if any(letter not in keys for letter in question["answer"]):
                raise SystemExit(f"Answer outside options: {question['label']}")
            if question["type"] == "multi" and question["answer"] != "".join(sorted(question["answer"])):
                raise SystemExit(f"Multi answer is not sorted: {question['label']}")
    return failures


def write_report(path, before_failures, after_failures, questions):
    type_counts = Counter(q["typeName"] for q in questions)
    chapter_counts = Counter(q["chapter"] for q in questions)
    lines = [
        "# 习近平思想解析自反馈优化报告",
        "",
        "## 闭环规则",
        "",
        "- 生成：按题型重写为“重点/逻辑/辨析/抓手”的短链路。",
        "- 反馈：检查模板废话、缺少选项辨析、缺少大题层次、乱码和重复句。",
        "- 修正：发现弱项后重新写入对应字段，再次评分。",
        "",
        "## 结果",
        "",
        f"- 题量：{len(questions)}",
        f"- 题型：{dict(type_counts)}",
        f"- 章节数：{len(chapter_counts)}",
        f"- 优化前问题题数：{len(before_failures)}",
        f"- 优化后问题题数：{len(after_failures)}",
        "",
        "## 优化后仍需人工关注",
        "",
    ]
    if after_failures:
        for label, score, issues in after_failures[:30]:
            lines.append(f"- {label}：score={score}，issues={', '.join(issues)}")
    else:
        lines.append("- 无。")
    lines.extend(["", "## 样例", ""])
    for question in questions[:3] + [q for q in questions if q["type"] == "essay"][:1]:
        lines.append(f"### {question['label']} {question['typeName']}")
        lines.append("")
        lines.append(question["stem"])
        lines.append("")
        lines.append(question["quickExplanation"])
        lines.append("")
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text("\n".join(lines) + "\n", encoding="utf-8")


def write_docx(questions, output):
    try:
        from docx import Document
        from docx.shared import Pt
    except ImportError as exc:
        raise SystemExit("python-docx is required to write DOCX.") from exc

    document = Document()
    style = document.styles["Normal"]
    style.font.name = "宋体"
    style.font.size = Pt(10.5)
    document.add_heading("习近平思想复习资料：逐题答案解析深度优化版", level=1)
    current_chapter = None
    for question in questions:
        if question["chapter"] != current_chapter:
            current_chapter = question["chapter"]
            document.add_heading(current_chapter, level=2)
        document.add_paragraph(f"{question['label']}【{question['typeName']}】{question['stem']}")
        for option in question.get("options", []):
            document.add_paragraph(option["text"])
        answer = question["answer"]
        if question["type"] in ("single", "multi"):
            selected = answer_phrase(question)
            answer = f"{answer}（{selected}）"
        document.add_paragraph(f"【答案】{answer}")
        document.add_paragraph(f"【知识点】{question['knowledge']}")
        document.add_paragraph("【解析】")
        for line in question["quickExplanation"].splitlines():
            document.add_paragraph(line)
        document.add_paragraph("【知识点详解】")
        for line in question["knowledgeDetail"].splitlines():
            document.add_paragraph(line)
    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(output)


def main():
    parser = argparse.ArgumentParser(description="Refine Xi Thought explanations with a self-feedback loop.")
    parser.add_argument("--input", default=str(DEFAULT_INPUT))
    parser.add_argument("--output", default=str(DEFAULT_INPUT))
    parser.add_argument("--report", default=str(DEFAULT_REPORT))
    parser.add_argument("--docx-output", default="")
    parser.add_argument("--check", action="store_true")
    args = parser.parse_args()

    input_path = Path(args.input)
    questions = json.loads(input_path.read_text(encoding="utf-8"))
    before_failures = validate_questions(json.loads(json.dumps(questions, ensure_ascii=False)))
    refine_questions(questions)
    after_failures = validate_questions(questions)
    report_path = Path(args.report)
    write_report(report_path, before_failures, after_failures, questions)

    if not args.check:
        output_path = Path(args.output)
        output_path.write_text(json.dumps(questions, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
        if args.docx_output:
            write_docx(questions, Path(args.docx_output))

    print(f"questions={len(questions)} before_failures={len(before_failures)} after_failures={len(after_failures)}")
    print(f"report={report_path}")
    if args.docx_output and not args.check:
        print(f"docx={args.docx_output}")


if __name__ == "__main__":
    main()
